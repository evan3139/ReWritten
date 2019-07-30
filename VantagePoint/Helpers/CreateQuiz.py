from docx import Document
import openpyxl

def create_quiz(excel, word, group):
    workbook = openpyxl.load_workbook(excel)
    ws = workbook.active

    # Date has to be the first row and column in order for this to work.
    date = ws.cell(row=1, column=1).value
    doc = Document()

    # Creates the top of the file with stuff that will not change.
    doc.add_paragraph(date + ":")
    doc.add_paragraph("Facilitator:")
    doc.add_paragraph("Topic:")
    doc.add_paragraph("Week:")
    doc.add_paragraph("Group:" + group)
    doc.add_paragraph("Questions:")
    doc.add_paragraph("")

    for i in range(2, ws.max_row):
        name = ws.cell(row=i, column=2).value
        age = ws.cell(row=i, column=6).value
        gender = ws.cell(row=i, column=7).value
        race = ws.cell(row=i, column=13).value
        doc.add_paragraph(str(gender) + ',' + str(age) + ',' + str(race) + ',' + str(name) + ':')
    doc.save(word)
