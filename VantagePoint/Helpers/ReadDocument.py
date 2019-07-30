from docx import Document
import os

from Helpers.Globals import HEADERS


def read_header(file):
    doc = Document(file)
    f = open("doc.txt", "w+")

    # Write the Docx to a txt then delete it. I dont like reading from a docx so I did this.
    for line in doc.paragraphs:
        line.text = line.text.strip()
        if line.text == "":
            continue
        f.write(line.text)
        f.write('\n')
    f.close()

    # Read Line by Line
    with open(f.name) as file:
        lines = file.readlines()

    # Split the lines at the ":" so we only keep what comes after
    lines = [x.partition(":")[0] for x in lines]
    lines = [x.strip() for x in lines]

    # Close and delete the .txt
    f.close()
    os.remove(f.name)

    return lines


def read_content(file):
    doc = Document(file)
    f = open("doc.txt", "w+")

    # Write the Docx to a txt then delete it. I dont like reading from a docx so I did this.
    for line in doc.paragraphs:
        line.text = line.text.strip()
        if line.text == "":
            continue
        f.write(line.text)
        f.write('\n')
    f.close()

    # Read Line by Line
    with open(f.name) as file:
        lines = file.readlines()

    # Make the headers Lowercase due to me making the Globals all lowercase you can change this if you wish
    headers = [x.partition(":")[0] for x in lines]
    headers = [x.strip() for x in headers]
    headers = [x.lower() for x in headers]

    # This pulls the right side of the ":" which are the contents essentially.
    content = [x.partition(":")[2] for x in lines]
    content = [x.strip() for x in content]

    # This will create a tuple of the header and content in a list.
    # Makes it much easier to associate headers with certain info BDI BAI etc.
    content_info = list(zip(headers, content))

    # Close and delete the .txt
    f.close()
    os.remove(f.name)

    for info in content_info:
        if HEADERS in info:
            print("yeet")
    return content_info
